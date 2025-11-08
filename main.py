from vosk import Model, KaldiRecognizer
import sounddevice as sd
import queue
import json
from docx import Document

# Model yolu 
model_path = r"vosk-model-small-tr-0.3"
model = Model(model_path)
recognizer = KaldiRecognizer(model, 16000)

# Word dosyası başlat 
doc = Document()
doc.add_heading("Asistan Ses Kaydı Metni", level=1)
doc.add_paragraph("Bu dosya Vosk modeliyle kaydedilmiştir.\n")
save_path = "asistan_kaydi.docx"

#  Ses kuyruğu 
audio_q = queue.Queue()

def save_text_to_word(text):
    """Her yeni cümlede Word dosyasına ekleme yapar."""
    doc.add_paragraph(text)
    doc.save(save_path)
    print(f" Kaydedildi: {text}")

def callback(indata, frames, time, status):
    """Ses verisi geldikçe kuyruğa ekler."""
    if status:
        print(status)
    audio_q.put(bytes(indata))

def listen_and_save():
    print(" Sistem aktif! Konuşmaya başlayabilirsin. (CTRL+C ile bitir)")
    try:
        with sd.InputStream(samplerate=16000, blocksize=8000, dtype='int16',
                            channels=1, callback=callback):
            while True:
                data = audio_q.get()
                if recognizer.AcceptWaveform(data):
                    result = json.loads(recognizer.Result())
                    text = result.get("text", "").strip()
                    if text:
                        print(" Tanınan Metin:", text)
                        save_text_to_word(text)

    except KeyboardInterrupt:
        print("\n Kayıt bitti. Dosya kaydediliyor...")
    finally:
        doc.save(save_path)
        print(f" Kayıt tamamlandı! Tüm konuşma '{save_path}' dosyasına kaydedildi.")

#  Program başlat
listen_and_save()
